"""
Document extraction MCP server for Claude Code.

Install into Claude Code:
  fastmcp install claude-code extract_mcp.py \
    --with pymupdf --with pdfplumber --with pytesseract \
    --with python-docx --with pillow --with pandas

Or manually:
  claude mcp add doc-extractor -- uv run \
    --with fastmcp --with pymupdf --with pdfplumber \
    --with pytesseract --with python-docx --with pillow --with pandas \
    fastmcp run extract_mcp.py
"""

import io
import base64
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
import pytesseract
import pandas as pd
import docx
from PIL import Image
from fastmcp import FastMCP

mcp = FastMCP(name="doc-extractor")


# ── helpers (same logic as your FastAPI app) ────────────────────────────────

def _extract_text_pdf(data: bytes) -> str:
    doc = fitz.open("pdf", data)
    return "\n".join(page.get_text("text") for page in doc)


def _extract_text_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def _extract_images_pdf(data: bytes) -> list[bytes]:
    doc = fitz.open("pdf", data)
    images = []
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            raw = doc.extract_image(xref)["image"]
            buf = io.BytesIO()
            Image.open(io.BytesIO(raw)).save(buf, format="PNG")
            images.append(buf.getvalue())
    return images


def _extract_images_docx(data: bytes) -> list[bytes]:
    document = docx.Document(io.BytesIO(data))
    images = []
    for rel in document.part.rels:
        if "image" in document.part.rels[rel].target_ref:
            raw = document.part.rels[rel].target_part.blob
            buf = io.BytesIO()
            Image.open(io.BytesIO(raw)).save(buf, format="PNG")
            images.append(buf.getvalue())
    return images


def _ocr(images: list[bytes]) -> str:
    return "\n".join(
        pytesseract.image_to_string(Image.open(io.BytesIO(img)))
        for img in images
    )


def _tables_pdf(data: bytes) -> list[dict]:
    tables = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            for i, table in enumerate(page.extract_tables(), 1):
                df = pd.DataFrame(table).dropna(how="all").dropna(axis=1, how="all")
                tables.append({"page": page_num, "table_index": i, "data": df.to_dict("records")})
    return tables


def _tables_docx(data: bytes) -> list[dict]:
    document = docx.Document(io.BytesIO(data))
    return [
        {
            "table_index": i,
            "data": [[cell.text.strip() for cell in row.cells] for row in table.rows],
        }
        for i, table in enumerate(document.tables, 1)
    ]


# ── MCP tools ────────────────────────────────────────────────────────────────

@mcp.tool()
def extract_from_file(file_path: str) -> dict:
    """
    Extract text, OCR text, table data, and image count from a PDF or DOCX file.

    Args:
        file_path: Absolute path to the PDF or DOCX file on disk.

    Returns:
        A dict with keys: file_type, extracted_text, ocr_text,
        extracted_tables (list of dicts), extracted_images (count).
    """
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        return {"error": f"File not found: {path}"}

    ext = path.suffix.lower().lstrip(".")
    if ext not in ("pdf", "docx"):
        return {"error": f"Unsupported file type '{ext}'. Use PDF or DOCX."}

    data = path.read_bytes()

    if ext == "pdf":
        text = _extract_text_pdf(data)
        images = _extract_images_pdf(data)
        tables = _tables_pdf(data)
    else:
        text = _extract_text_docx(data)
        images = _extract_images_docx(data)
        tables = _tables_docx(data)

    ocr_text = _ocr(images) if images else ""

    return {
        "file_type": ext.upper(),
        "extracted_text": text,
        "ocr_text": ocr_text,
        "extracted_images": len(images),
        "extracted_tables": tables,
    }


@mcp.tool()
def extract_text_only(file_path: str) -> str:
    """
    Fast path: extract only the selectable text from a PDF or DOCX.
    Skips OCR and table extraction. Use this when you just need the body text.

    Args:
        file_path: Absolute path to the file.
    """
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        return f"File not found: {path}"

    ext = path.suffix.lower().lstrip(".")
    data = path.read_bytes()

    if ext == "pdf":
        return _extract_text_pdf(data)
    elif ext == "docx":
        return _extract_text_docx(data)
    else:
        return f"Unsupported file type: {ext}"


@mcp.tool()
def extract_tables_only(file_path: str) -> list[dict]:
    """
    Extract only the tables from a PDF or DOCX file.

    Args:
        file_path: Absolute path to the file.

    Returns:
        List of table dicts. PDF tables include a 'page' key; DOCX tables do not.
    """
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        return [{"error": f"File not found: {path}"}]

    ext = path.suffix.lower().lstrip(".")
    data = path.read_bytes()

    if ext == "pdf":
        return _tables_pdf(data)
    elif ext == "docx":
        return _tables_docx(data)
    else:
        return [{"error": f"Unsupported file type: {ext}"}]


if __name__ == "__main__":
    mcp.run(transport="stdio")
